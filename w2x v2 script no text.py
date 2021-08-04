import docx;
import openpyxl;
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from tkinter import *
from tkinter import filedialog

tabelNum = 0
systemRow = 0
sectionRow = 1
word = "wordNull"
excel = "excelNull"
    
def openWord():
    filepath = filedialog.askopenfilename()
    global word
    word = filepath
    
def openExcel():
    filepath = filedialog.askopenfilename()
    global excel
    excel = filepath

def getV ():  
    global tabelNum
    global systemRow
    global sectionRow
    x1 = (entry1.get())
    tabelNum = int(x1)-1
    x2 = (entry2.get())
    systemRow = int(x2)-1
    x3 = (entry3.get())
    sectionRow = int(x3)-1
    window.destroy()
      
import tkinter as tk
window = Tk()
window.geometry("450x330")

l = Label(window, text="")
l.pack()
l.config(font=('ariel' ,30, 'bold'))

l2 = Label(window, text="")
l2.pack()
l2.config(font=('ariel' ,8, 'bold'))
    
entry1 = tk.Entry (window)
entry1.pack(side=TOP)
x1 = entry1.get()

l5 = Label(window, text="")
l5.pack()
l5.config(font=('ariel' ,8, 'bold'))

entry3 = tk.Entry (window)
entry3.pack(side=TOP)
x3 = entry3.get()

l4 = Label(window, text="")
l4.pack()
l4.config(font=('ariel' ,8, 'bold'))

entry2 = tk.Entry (window)
entry2.pack(side=TOP)
x2 = entry2.get()

button = Button(window, text="word", command=openWord)
button.pack(side=LEFT)
button.config(font=('ariel' ,40, 'bold'))
button.config(bg='#4284f5')
print(word)
    
button2 = Button(window, text="excel", command=openExcel)
button2.pack(side=RIGHT)
button2.config(font=('ariel' ,40, 'bold'))
button2.config(bg='#2ea63c')
print(excel)

button3 = Button(window, text="", height = 1, width = 100, command=getV)
button3.pack(side=BOTTOM)
button3.config(font=('ariel' ,15, 'bold'))
button3.config(bg='#e3f542')
    
window.mainloop()
filearry = [word,excel]
print(filearry)
print("wordPath : ", filearry[0])
print("excelPath : ", filearry[1])
print("tabelNum : ", tabelNum, "(input : ", tabelNum+1, ")")
print("systemRow : ", systemRow, "(input : ", systemRow+1, ")")
print("sectionRow : ", sectionRow, "(input : ", sectionRow+1, ")")

def lastValuableCell():
    count2 = 1
    var = ws['A' + str(count2)].value
    flag = True

    while flag == True:
        var = ws['A' + str(count2)].value
        if var is None:
            print("var has a value of None")
            flag = False
            print (flag)
        else:
            flag = True
            print("var has a value")
            print(ws['A' + str(count2)].value)
            print(flag)
            count2 += 1
    lastcell = count2
    print(lastcell)
    return lastcell

doc = docx.Document(filearry[0])
celltext = doc.tables[tabelNum].cell(0,0).text
cellname = doc.tables[tabelNum].cell(0,1).text

print(cellname + ':')
#print(celltext)

textArry = celltext.splitlines()
print(textArry)


wb = load_workbook(filearry[1])
ws = wb.active

num2 = 0
while num2 <= len(doc.tables[tabelNum].rows)-1:
    celltext = doc.tables[tabelNum].cell(num2, systemRow).text
    cellname = doc.tables[tabelNum].cell(num2, sectionRow).text
    textArry = celltext.splitlines()
    x = lastValuableCell() 
    for i in textArry:
        print(i)
        print (x)
        ws['A' + str(x)] = i
        ws['B' + str(x)] = cellname
        x += 1
    num2 += 1
wb.save(filearry[1])
